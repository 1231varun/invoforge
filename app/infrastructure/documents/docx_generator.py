"""DOCX Document Generator Implementation"""
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from app.core.entities.invoice import Invoice
from app.core.entities.settings import Settings
from app.core.interfaces.document_generator import DocumentGenerator

PROJECT_ROOT = Path(__file__).parent.parent.parent.parent


class DocxGenerator(DocumentGenerator):
    """
    Generates DOCX invoices using python-docx.
    
    Creates professional export invoices with proper Calibri formatting,
    consistent spacing, and styled tables.
    """

    FONT_NAME = "Calibri"
    FONT_SIZE = Pt(11)
    SPACING_AFTER = Pt(2)
    SECTION_SPACING = Pt(14)

    def __init__(self, output_dir: Path = None):
        self._output_dir = output_dir or PROJECT_ROOT / "output"
        self._output_dir.mkdir(exist_ok=True)

    def generate(self, invoice: Invoice, settings: Settings) -> Path:
        """Generate DOCX document and return file path"""
        doc = Document()
        self._configure_document(doc)

        self._add_heading(doc)
        self._add_invoice_details(doc, invoice)
        self._add_supplier_details(doc, settings)
        self._add_export_details(doc, invoice, settings)
        self._add_client_details(doc, settings)
        self._add_working_days(doc, invoice)
        self._add_services_table(doc, invoice, settings)
        self._add_totals(doc, invoice, settings)
        self._add_bank_details(doc, settings)
        self._add_declaration(doc)
        self._add_signature(doc, settings)

        filepath = self._get_filepath(invoice, settings)
        doc.save(str(filepath))

        return filepath

    def _configure_document(self, doc: Document):
        """Configure document styles and margins for single-page invoice"""
        # Set default font for entire document
        style = doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE
        style.paragraph_format.space_after = self.SPACING_AFTER
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Ensure font applies to all character types
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.FONT_NAME)
        rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
        rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        rFonts.set(qn('w:cs'), self.FONT_NAME)

        # Set compact margins for single-page layout
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

    def _set_font(self, run, size: Optional[Pt] = None, bold: bool = False):
        """Apply Calibri font to a run with proper XML handling"""
        run.font.name = self.FONT_NAME
        run.font.size = size or self.FONT_SIZE
        run.bold = bold

        # Ensure font applies to all character types in XML
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.FONT_NAME)
        rFonts.set(qn('w:hAnsi'), self.FONT_NAME)
        rFonts.set(qn('w:eastAsia'), self.FONT_NAME)
        rFonts.set(qn('w:cs'), self.FONT_NAME)

    def _set_paragraph_spacing(self, paragraph, space_after: Optional[Pt] = None, space_before: Optional[Pt] = None):
        """Set paragraph spacing"""
        pf = paragraph.paragraph_format
        pf.space_after = space_after if space_after is not None else self.SPACING_AFTER
        pf.space_before = space_before if space_before is not None else Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False,
                       size: Optional[Pt] = None, space_after: Optional[Pt] = None) -> object:
        """Add a formatted paragraph"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        self._set_font(run, size=size, bold=bold)
        self._set_paragraph_spacing(p, space_after=space_after)
        return p

    def _add_heading(self, doc: Document):
        """Add centered underlined heading"""
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("EXPORT INVOICE (UNDER LUT)")
        self._set_font(run, size=Pt(13), bold=True)
        run.underline = True
        self._set_paragraph_spacing(heading, space_after=Pt(12))

    def _add_invoice_details(self, doc: Document, invoice: Invoice):
        """Add invoice number, date, and service period"""
        self._add_paragraph(doc, f"Invoice No.: {invoice.invoice_number}")
        self._add_paragraph(doc, f"Invoice Date: {invoice.invoice_date.isoformat()}")
        self._add_paragraph(
            doc,
            f"Service Period: {invoice.service_period_start.isoformat()} "
            f"to {invoice.service_period_end.isoformat()}",
            space_after=self.SECTION_SPACING
        )

    def _add_supplier_details(self, doc: Document, settings: Settings):
        """Add supplier business details"""
        self._add_paragraph(doc, settings.supplier_name)
        for line in settings.supplier_address.replace("\\n", "\n").split("\n"):
            if line.strip():
                self._add_paragraph(doc, line)
        self._add_paragraph(doc, f"GSTIN: {settings.gstin}", space_after=Pt(1))
        self._add_paragraph(doc, f"PAN: {settings.pan}", space_after=Pt(1))
        self._add_paragraph(doc, f"Email: {settings.supplier_email}", space_after=self.SECTION_SPACING)

    def _add_export_details(self, doc: Document, invoice: Invoice, settings: Settings):
        """Add export and LUT details"""
        self._add_paragraph(doc, f"Place of Supply: {settings.place_of_supply}")
        self._add_paragraph(doc, "Type: Export of Services without Payment of IGST (Under LUT)")
        self._add_paragraph(
            doc,
            f"LUT No.: {settings.lut_no}     Validity: FY {invoice.validity_year}",
            space_after=self.SECTION_SPACING
        )

    def _add_client_details(self, doc: Document, settings: Settings):
        """Add client billing details"""
        self._add_paragraph(doc, "Bill To:", bold=True)
        self._add_paragraph(doc, f"Name:  {settings.client_name}")
        self._add_paragraph(doc, f"Address:  {settings.client_address}")
        self._add_paragraph(
            doc,
            f"Country: {settings.client_country}    Email: {settings.client_email}",
            space_after=self.SECTION_SPACING
        )

    def _add_working_days(self, doc: Document, invoice: Invoice):
        """Add working days and leaves section"""
        # Total working days
        p = doc.add_paragraph()
        run1 = p.add_run("Total No. of working days: ")
        self._set_font(run1, bold=True)
        run2 = p.add_run(str(invoice.total_working_days))
        self._set_font(run2)
        self._set_paragraph_spacing(p)

        # Leaves taken
        p2 = doc.add_paragraph()
        run3 = p2.add_run("Leaves Taken: ")
        self._set_font(run3, bold=True)
        run4 = p2.add_run(str(invoice.leaves_taken))
        self._set_font(run4)
        self._set_paragraph_spacing(p2)

        # Leave dates
        for leave_date in invoice.leave_dates:
            self._add_paragraph(doc, leave_date.isoformat())

        # Section spacer
        self._add_paragraph(doc, "", space_after=self.SECTION_SPACING)

    def _set_table_borders(self, table):
        """Set explicit black borders on table"""
        tbl = table._tbl

        # Find or create tblPr element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove existing borders if any
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)

        # Create new borders
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 1/2 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _set_cell_padding(self, cell, top: int = 60, bottom: int = 60,
                          left: int = 80, right: int = 80):
        """Set cell padding in twips"""
        tc = cell._tc

        # Find or create tcPr element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)

        # Remove existing margins if any
        existing_mar = tcPr.find(qn('w:tcMar'))
        if existing_mar is not None:
            tcPr.remove(existing_mar)

        # Create new margins
        tcMar = OxmlElement('w:tcMar')

        for margin_name, margin_value in [('top', top), ('bottom', bottom),
                                           ('left', left), ('right', right)]:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), str(margin_value))
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)

        tcPr.append(tcMar)

    def _set_cell_content(self, cell, text: str, bold: bool = False):
        """Set cell content with proper font"""
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        self._set_font(run, bold=bold)
        self._set_paragraph_spacing(p, space_after=Pt(0), space_before=Pt(0))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _add_services_table(self, doc: Document, invoice: Invoice, settings: Settings):
        """Add services description table with proper formatting"""
        # Spacer before table
        spacer = doc.add_paragraph()
        self._set_paragraph_spacing(spacer, space_after=Pt(6))

        # Create table
        table = doc.add_table(rows=2, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        # Set explicit borders
        self._set_table_borders(table)

        # Set table width
        table.columns[0].width = Inches(6.3)

        # Header row
        header_cell = table.rows[0].cells[0]
        self._set_cell_content(header_cell, "Description of Services:", bold=True)
        self._set_cell_padding(header_cell)

        # Content row
        content_cell = table.rows[1].cells[0]
        content_text = (
            f"1. {settings.service_description} | "
            f"Days: {invoice.days_worked} | "
            f"Rate: {settings.currency} {invoice.rate:.2f} | "
            f"Amount: {settings.currency} {invoice.amount:.2f}"
        )
        self._set_cell_content(content_cell, content_text)
        self._set_cell_padding(content_cell)

        # Spacer after table
        spacer2 = doc.add_paragraph()
        self._set_paragraph_spacing(spacer2, space_after=Pt(6))

    def _add_totals(self, doc: Document, invoice: Invoice, settings: Settings):
        """Add totals section"""
        self._add_paragraph(doc, "Tax: Export Without Payment of IGST under LUT")
        self._add_paragraph(doc, f"Total Payable ({settings.currency}): {invoice.total_payable:.2f}")
        self._add_paragraph(
            doc,
            f"Total in Words: {invoice.amount_in_words}",
            space_after=self.SECTION_SPACING
        )

    def _add_bank_details(self, doc: Document, settings: Settings):
        """Add bank details for wire transfer"""
        self._add_paragraph(
            doc,
            f"Bank Details: {settings.account_holder} | "
            f"A/c No: {settings.account_no} | "
            f"Bank: {settings.bank_name} | "
            f"SWIFT: {settings.swift_code}",
            space_after=self.SECTION_SPACING
        )

    def _add_declaration(self, doc: Document):
        """Add export declaration"""
        self._add_paragraph(
            doc,
            "Declaration: This invoice is issued as Export of Services without payment "
            "of IGST under LUT as per Sec. 16 of IGST Act. Payment will be received in "
            "foreign currency. No GST charged.",
            space_after=self.SECTION_SPACING
        )

    def _add_signature(self, doc: Document, settings: Settings):
        """Add signature section"""
        self._add_paragraph(doc, f"For {settings.account_holder}")
        # Use signatory_name if set, otherwise fall back to supplier_name
        signatory = settings.signatory_name
        if not signatory and settings.supplier_name:
            # Extract name without parenthetical as fallback
            signatory = settings.supplier_name.split('(')[0].strip()
        if signatory:
            self._add_paragraph(doc, signatory)

    def _get_filepath(self, invoice: Invoice, settings: Settings) -> Path:
        """Generate output filepath"""
        month_name = invoice.invoice_date.strftime("%B")
        year = invoice.invoice_date.year
        client = settings.client_name.split()[0] if settings.client_name else "Invoice"
        filename = f"{client} Invoice - {invoice.invoice_number} - {month_name} {year}.docx"
        return self._output_dir / filename
